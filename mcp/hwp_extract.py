#!/usr/bin/env python3
"""HWP 5.0 텍스트 추출 (스펙 준수 개선판) — 부패영향평가 AI MCP용

DataMan extract_hwp_direct 대비 개선점:
  1. 인라인 컨트롤(4~9, 19, 20)을 16바이트로 정확히 소비
     — 기존엔 탭(9)을 2바이트 처리해 조 제목 단락이 통째로 깨짐(汫ॣ)
  2. UTF-16 서로게이트(한컴 특수문자: 체크박스 등) → '□' 대체
     — pyhwp는 이 지점에서 ParseError로 중단
  3. 표 안 텍스트도 동일 단락 스트림에서 수집 (별표 체크리스트 보존)
"""
import struct
import zlib

import olefile

HWPTAG_PARA_TEXT = 0x10 + 51  # 67

# HWP 5.0 문자 컨트롤 분류 (한글문서파일형식 5.0 스펙)
CHAR_CONTROLS = {0, 10, 13, 24, 25, 26, 27, 28, 29, 30, 31}      # 2바이트
SKIP16_CONTROLS = {1, 2, 3, 4, 5, 6, 7, 8, 11, 12, 14, 15, 16,    # 16바이트
                   17, 18, 19, 20, 21, 22, 23}                     # (인라인+확장)
TAB = 9                                                            # 16바이트 + '\t'


def _parse_para_text(data: bytes) -> str:
    chars = []
    i = 0
    n = len(data)
    while i + 2 <= n:
        code = struct.unpack_from("<H", data, i)[0]
        if code == TAB:
            chars.append("\t")
            i += 16
            continue
        if code in SKIP16_CONTROLS:
            i += 16
            continue
        i += 2
        if code in CHAR_CONTROLS:
            if code in (10, 13):
                chars.append("\n")
            continue
        if 0xD800 <= code <= 0xDFFF:
            # 정상 서로게이트 쌍이면 결합, 한컴 단독 서로게이트(특수문자)면 □
            if code <= 0xDBFF and i + 2 <= n:
                lo = struct.unpack_from("<H", data, i)[0]
                if 0xDC00 <= lo <= 0xDFFF:
                    i += 2
                    chars.append(chr(0x10000 + ((code - 0xD800) << 10) + (lo - 0xDC00)))
                    continue
            chars.append("□")
            continue
        chars.append(chr(code))
    return "".join(chars)


def extract_hwp_text(path: str) -> str:
    """HWP 파일 → 본문 전체 텍스트 (문단 단위 개행)"""
    ole = olefile.OleFileIO(path)
    try:
        header = ole.openstream("FileHeader").read()
        is_compressed = bool(header[36] & 1)
        paragraphs = []
        section_idx = 0
        while True:
            name = f"BodyText/Section{section_idx}"
            if not ole.exists(name):
                break
            raw = ole.openstream(name).read()
            if is_compressed:
                try:
                    raw = zlib.decompress(raw, -15)
                except zlib.error:
                    section_idx += 1
                    continue
            pos = 0
            while pos + 4 <= len(raw):
                hdr = struct.unpack_from("<I", raw, pos)[0]
                tag_id = hdr & 0x3FF
                size = (hdr >> 20) & 0xFFF
                pos += 4
                if size == 0xFFF:
                    if pos + 4 > len(raw):
                        break
                    size = struct.unpack_from("<I", raw, pos)[0]
                    pos += 4
                if pos + size > len(raw):
                    break
                if tag_id == HWPTAG_PARA_TEXT:
                    text = _parse_para_text(raw[pos:pos + size]).strip()
                    if text:
                        paragraphs.append(text)
                pos += size
            section_idx += 1
        return "\n".join(paragraphs)
    finally:
        ole.close()


if __name__ == "__main__":
    import sys
    print(extract_hwp_text(sys.argv[1]))


def extract_hwpx_text(path: str) -> str:
    """HWPX(OWPML zip) → 텍스트 (XML 엔티티 해제 포함)"""
    import html
    import re
    import zipfile
    texts = []
    with zipfile.ZipFile(path) as z:
        sections = sorted(n for n in z.namelist()
                          if re.match(r"Contents/section\d+\.xml", n))
        for name in sections:
            xml = z.read(name).decode("utf-8", errors="ignore")
            xml = re.sub(r"</hp:p>", "\n", xml)
            texts.append(html.unescape(re.sub(r"<[^>]+>", "", xml)))
    return "\n".join(texts)


def extract_pdf_text(path: str) -> str:
    """PDF → 텍스트 (PyMuPDF) — 타 기관이 규정을 PDF로 공시하는 경우 대응"""
    import fitz
    doc = fitz.open(path)
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def extract_docx_text(path: str) -> str:
    """DOCX → 텍스트 (표 포함)"""
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_hwpml_text(path: str) -> str:
    """HWPML(한글 XML, <HWPML> 루트) → 텍스트. BODY의 <P> 단락 단위 개행.
    구형 .hwp로 표기됐으나 실제 HWPML(2.x)로 공시된 문서 대응(회계기준 등)."""
    import html
    import re
    with open(path, encoding="utf-8", errors="ignore") as f:
        xml = f.read()
    m = re.search(r"<BODY\b.*</BODY>", xml, re.DOTALL)   # 본문만 (HEAD 메타 제외)
    body = m.group(0) if m else xml
    body = re.sub(r"</P\s*>", "\n", body)                 # 단락 경계 → 개행
    body = re.sub(r"<[^>]+>", "", body)                   # 잔여 태그 제거
    return re.sub(r"\n{3,}", "\n\n", html.unescape(body)).strip()


def extract_any(path: str) -> str:
    """확장자에 따라 HWP/HWPX/PDF/DOCX 자동 분기 추출"""
    import os
    ext = os.path.splitext(path)[1].lower()
    if ext == ".hwpx":
        return extract_hwpx_text(path)
    if ext == ".pdf":
        return extract_pdf_text(path)
    if ext == ".docx":
        return extract_docx_text(path)
    return extract_hwp_text(path)
